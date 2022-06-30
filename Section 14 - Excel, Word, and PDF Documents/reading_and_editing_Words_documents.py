import docx

print('\n1------------------------------------------------\n')
# .Document()
# .paragraphs

d = docx.Document('demo.docx')
print(d.paragraphs)

print(d.paragraphs[0]) # <docx.text.paragraph.Paragraph object at 0x03AEF370>

print(d.paragraphs[0].text) # Document Title - the first paragraph is 'Document Title'

print(d.paragraphs[1].text) # A plain paragraph having some bold and some italic.

print('\n2------------------------------------------------\n')
# .runs

p = d.paragraphs[1]
print(p.runs) # every bold or italic will start with a new runs

print(p.runs[0].text) # A plain paragraph having some - first runs

print(p.runs[1].text) # bold
print(p.runs[1].bold) # True

print(p.runs[2].text) #  and some 

print(p.runs[3].text) # italic.
print(p.runs[3].italic) # True

print('\n3------------------------------------------------\n')
# edit and save the docx document
# .save()

p.runs[3].underline = True # set the 'italic' word with underline
p.runs[3].text = 'italic and underline.' # change the word 'italic' to 'italic and underline.'

d.save('demo2.docx') # the new docx file was saved with name 'demo2.docx'

print('\n4------------------------------------------------\n')
# .style

print(p.style) # _ParagraphStyle('Normal') id: 71393488 - Normal style

p.style = 'Title' # change 'A plain paragraph having some bold and some italic.' to 'Title' style
d.save('demo3.docx') 

print('\n5------------------------------------------------\n')
# create a new docx file
# .add_paragraph()

d = docx.Document()
d.add_paragraph('hello this is a paragraph.')
d.add_paragraph('This is another paragraph.')

d.save('demo4.docx')

print('\n6------------------------------------------------\n')
# .add_runs

p = d.paragraphs[0]
p.add_run('This is a new run.')
p.runs[1].bold = True
d.save('demo5.docx')




























